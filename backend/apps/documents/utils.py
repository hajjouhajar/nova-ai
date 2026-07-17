import io
import logging

import pdfplumber
from pypdf import PdfReader

try:
    import pypdfium2 as pdfium
    from PIL import Image
    import pytesseract
    PDFIUM_AVAILABLE = True
except ImportError:
    PDFIUM_AVAILABLE = False

try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

import docx

from apps.ai.ollama import embed

logger = logging.getLogger(__name__)

# Chemins Tesseract possibles sur Windows
TESSERACT_PATHS = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    r"C:\Users\user\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
]


def _configure_tesseract() -> bool:
    """Cherche et configure tesseract.exe automatiquement."""
    import os
    try:
        import pytesseract
        for path in TESSERACT_PATHS:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                return True
    except ImportError:
        pass
    return False


def _ocr_via_pdfium(data: bytes) -> str:
    """
    OCR via pypdfium2 (déjà installé) + Tesseract.
    pypdfium2 rend chaque page en image haute résolution,
    puis Tesseract en extrait le texte.
    """
    import pytesseract
    if not _configure_tesseract():
        raise ValueError(
            "Ce PDF est scanné (images uniquement). "
            "Installez Tesseract OCR pour lire ce type de document : "
            "https://github.com/UB-Mannheim/tesseract/wiki"
        )

    pdf = pdfium.PdfDocument(data)
    texte_pages = []
    for page_index in range(len(pdf)):
        page = pdf[page_index]
        # Rendu à 300 DPI (scale=300/72)
        bitmap = page.render(scale=300 / 72, rotation=0)
        img = bitmap.to_pil()
        # OCR français + anglais
        t = pytesseract.image_to_string(img, lang="fra+eng", config="--psm 6")
        if t.strip():
            texte_pages.append(t.strip())
    pdf.close()
    return "\n\n".join(texte_pages)


def _extraire_texte_pdf(fichier) -> str:
    """
    Extrait le texte d'un PDF avec 3 niveaux de fallback :
      1. pdfplumber  — meilleur pour PDFs complexes (tableaux, colonnes)
      2. pypdf       — fallback léger pour PDFs simples
      3. OCR         — pour les PDFs scannés (pypdfium2 + Tesseract)
    """
    data = fichier.read() if hasattr(fichier, "read") else fichier

    # ── Niveau 1 : pdfplumber ────────────────────────────────────────────────
    try:
        pages = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t and t.strip():
                    pages.append(t.strip())
        texte = "\n\n".join(pages)
        if len(texte.strip()) >= 50:
            logger.info("PDF extrait via pdfplumber (%d caractères)", len(texte))
            return texte
    except Exception as exc:
        logger.warning("pdfplumber a échoué : %s", exc)

    # ── Niveau 2 : pypdf ─────────────────────────────────────────────────────
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t and t.strip():
                pages.append(t.strip())
        texte = "\n\n".join(pages)
        if len(texte.strip()) >= 50:
            logger.info("PDF extrait via pypdf (%d caractères)", len(texte))
            return texte
    except Exception as exc:
        logger.warning("pypdf a échoué : %s", exc)

    # ── Niveau 3 : OCR (pypdfium2 + Tesseract) ───────────────────────────────
    if not PDFIUM_AVAILABLE:
        raise ValueError(
            "Ce PDF est scanné (images uniquement) et les librairies OCR "
            "ne sont pas disponibles."
        )

    logger.info("PDF scanné détecté — tentative OCR...")
    texte = _ocr_via_pdfium(data)
    if len(texte.strip()) >= 20:
        logger.info("PDF extrait via OCR (%d caractères)", len(texte))
        return texte

    raise ValueError(
        "Ce PDF est entièrement scanné et Tesseract OCR n'est pas encore installé. "
        "Pour lire ce type de document, installez Tesseract : "
        "https://github.com/UB-Mannheim/tesseract/wiki — "
        "puis redémarrez le serveur."
    )


def extraire_texte(fichier, type_fichier: str) -> str:
    """Point d'entrée principal pour l'extraction de texte."""
    if type_fichier == "pdf":
        return _extraire_texte_pdf(fichier)
    if type_fichier == "docx":
        document = docx.Document(fichier)
        return "\n".join(p.text for p in document.paragraphs if p.text.strip())
    # txt / md — lecture directe
    if hasattr(fichier, "read"):
        return fichier.read().decode("utf-8", errors="ignore")
    return fichier.decode("utf-8", errors="ignore")


def decouper_en_chunks(texte: str, taille_mots: int = 300, chevauchement: int = 50) -> list[str]:
    mots = texte.split()
    return [
        " ".join(mots[i: i + taille_mots])
        for i in range(0, len(mots), taille_mots - chevauchement)
        if mots[i: i + taille_mots]
    ]


def generer_embedding(texte: str) -> list[float]:
    return embed(texte)
